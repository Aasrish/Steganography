# -*- coding: utf-8 -*-
"""B20CS042.ipynb

Automatically generated by Colaboratory.

Original file is located at
    https://colab.research.google.com/drive/1AS_EMorojmwCjbQ1A3TOmEzrbzfbvGIO

# ***Semagram-based steganography***
"""

!pip install pillow
from PIL import Image

import numpy as np
from sklearn.cluster import KMeans
image = Image.open("/content/Knight.png")
pixel_values = np.array(image)
pixel_values = pixel_values.reshape((-1, 3))
kmeans = KMeans(n_clusters=16, random_state=0).fit(pixel_values)
new_pixel_values = kmeans.cluster_centers_[kmeans.labels_]
new_pixel_values = new_pixel_values.reshape(image.size + (3,))
new_image = Image.fromarray(np.uint8(new_pixel_values))
new_image.save("reduced_image.png")
image2 = Image.open('reduced_image.png')
image2 = image2.convert('RGB')
width, height = image2.size

colors2 = []
for x in range(height):
  for y in range(width):
    pixel = image2.getpixel((y, x))
    colors2.append(pixel)

print(width)
print(height)
print(len(colors2))

!pip install python-docx

from docx import Document
from docx.shared import RGBColor
document = Document()
document.add_paragraph('')
document.save('colored_text.docx')

i = 0
text = "10"
while i < len(colors2):
  j = i%360
  k = i//360
  colr = colors2[i]
  if j == 0:
    para = document.add_paragraph()
    run = para.add_run(text)
    font = run.font
    font.color.rgb = RGBColor(colr[0], colr[1], colr[2])
    print(k)
  else:
    last_para = document.paragraphs[-1]
    run = last_para.add_run(text)
    font = run.font
    font.color.rgb = RGBColor(colr[0], colr[1], colr[2])
  i=i+1
document.save('/content/colored_text.docx')
